"""
Generiert das Arbeitsblatt zur AGB-Präsentation.
Zwei Versionen: Schüler-AB (leer) und Lehrer-Lösungsblatt.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


WA_GREEN = RGBColor(0x16, 0x9C, 0x4C)
SNAP_YELLOW = RGBColor(0xC9, 0xA8, 0x00)
COMPARE_CYAN = RGBColor(0x05, 0x84, 0x96)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x66, 0x66, 0x66)


def set_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_heading(doc, text, size=16, color=DARK, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_para(doc, text, size=11, bold=False, italic=False, color=DARK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_lines(doc, count=2):
    """Schreiblinien hinzufügen — leere Zeilen mit unterstrichenem Abstand."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Unterstrich auf 14 cm
        run = p.add_run("_" * 90)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_kopfzeile(doc, titel, untertitel):
    """Logo-freie Kopfzeile mit Titel + Name/Klasse/Datum-Zeilen."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(titel)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK
    r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(untertitel)
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = GREY

    # Name | Klasse | Datum
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [Cm(7), Cm(4), Cm(5)]
    for i, w in enumerate(widths):
        t.columns[i].width = w
    cells = t.rows[0].cells
    for i, label in enumerate(["Name: _____________________________",
                                "Klasse: __________",
                                "Datum: ___________"]):
        cells[i].width = widths[i]
        para = cells[i].paragraphs[0]
        run = para.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
    # Abstand nach der Tabelle
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def aufgabe_header(doc, nr, titel, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"Aufgabe {nr}  ·  ")
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = color
    r1.font.name = "Calibri"
    r2 = p.add_run(titel)
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = DARK
    r2.font.name = "Calibri"


def baue_arbeitsblatt(loesungen=False):
    doc = Document()
    set_margins(doc)

    # Standard-Schrift
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if loesungen:
        add_kopfzeile(doc,
            "🔓  Lösungsblatt: Was du wirklich zustimmst",
            "WhatsApp & Snapchat AGB · Lehrer-Version mit Erwartungshorizont")
    else:
        add_kopfzeile(doc,
            "Was du wirklich zustimmst",
            "WhatsApp & Snapchat AGB · Begleitheft zur Präsentation")

    # ──────────────────────────────────────────────────────────
    # AUFGABE 1 — Schätzen vor der Präsentation
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 1, "Schätzen — bevor die Präsentation startet",
                   color=COMPARE_CYAN)
    add_para(doc, "Wie schätzt du es ein? Kreuze an oder schreibe deine Vermutung. Wir lösen es danach gemeinsam auf.",
             italic=True, color=GREY, space_after=10)

    # 5 Items
    items = [
        ("Ab welchem Alter darf man laut WhatsApp-AGB die App nutzen?",
         "☐ 10 Jahre  ☐ 13 Jahre  ☐ 16 Jahre  ☐ 18 Jahre",
         "13 Jahre (laut AGB) — aber DSGVO sagt: ohne Eltern erst ab 16."),
        ("Wenn du eine Snapchat-Story löschst, sind die Daten weg — stimmt das?",
         "☐ Ja, sofort gelöscht.   ☐ Nach 24 Stunden weg.   ☐ Nein.",
         "Nein. Snap darf Server-Kopien öffentlicher Inhalte „auf unbestimmte Zeit\" behalten."),
        ("Bekommst du Geld, wenn Snapchat dein Spotlight-Video kommerziell nutzt?",
         "☐ Ja, anteilig.   ☐ Nein, gar nichts.   ☐ Weiß ich nicht.",
         "Nein. Die Lizenz ist „gebührenfrei und unbefristet\". Du bekommst 0 €."),
        ("Lädt WhatsApp dein Adressbuch hoch — auch von Leuten ohne WhatsApp?",
         "☐ Nein, nur WhatsApp-Nutzer.   ☐ Ja, alle Nummern.   ☐ Nur wenn ich erlaube.",
         "Ja, alle Nummern werden hochgeladen, sobald du Kontakte freigibst."),
        ("Was ist Snaps maximale Haftung, wenn dir durch ihre Schuld Schaden entsteht?",
         "☐ unbegrenzt   ☐ 1 Mio. USD   ☐ 10.000 USD   ☐ 100 USD",
         "100 US-Dollar — egal wie groß der Schaden ist."),
    ]
    for q, opts, ans in items:
        # Frage
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"  ▸  {q}")
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = DARK
        # Optionen
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(opts)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GREY
        if loesungen:
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(0.5)
            p3.paragraph_format.space_after = Pt(6)
            r3 = p3.add_run("✓  " + ans)
            r3.font.size = Pt(10)
            r3.font.bold = True
            r3.font.color.rgb = WA_GREEN

    # ──────────────────────────────────────────────────────────
    # AUFGABE 2 — Notizen während der Präsentation
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 2, "Mitschreiben während der Präsentation",
                   color=WA_GREEN)
    add_para(doc, "Notiere die Aha-Momente, die dich am meisten überraschen. Pro App mindestens 3 Punkte.",
             italic=True, color=GREY, space_after=8)

    add_para(doc, "💬  WhatsApp:", bold=True, color=WA_GREEN, size=12, space_after=4)
    if loesungen:
        for line in [
            "• Mindestalter laut AGB nur 13 — DSGVO sagt eigentlich 16.",
            "• Adressbuch wird komplett hochgeladen, auch Nicht-Nutzer.",
            "• Lizenz an Inhalten: weltweit, gebührenfrei, übertragbar (begrenzt auf Service).",
            "• Meta-AI-Chats heben die E2E-Verschlüsselung auf.",
            "• Account löschen ≠ Daten weg: Lizenzen bleiben, 30 Tage Server-Speicherung.",
        ]:
            add_para(doc, line, size=10, color=DARK, space_after=2)
    else:
        add_lines(doc, 5)

    add_para(doc, "👻  Snapchat:", bold=True, color=SNAP_YELLOW, size=12, space_after=4)
    if loesungen:
        for line in [
            "• Mindestalter 13 (DSGVO 16); bei Verstoß Account-Löschung.",
            "• Öffentliche Inhalte: unwiderrufliche, unbefristete, kommerzielle Lizenz — kein Cent für dich.",
            "• KEINE E2E-Verschlüsselung — Snap kann jederzeit mitlesen.",
            "• Server-Kopien öffentlicher Inhalte „auf unbestimmte Zeit\".",
            "• My AI: alle Inputs werden zu lizenzpflichtigen „Inhalten\".",
            "• Strafverfolgung erhält Daten; Klagen nur via US-Schiedsverfahren; max. Haftung 100 USD.",
        ]:
            add_para(doc, line, size=10, color=DARK, space_after=2)
    else:
        add_lines(doc, 5)

    # ──────────────────────────────────────────────────────────
    # AUFGABE 3 — Vergleichstabelle ausfüllen
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 3, "Vergleich: WhatsApp vs. Snapchat",
                   color=COMPARE_CYAN)
    add_para(doc, "Fülle die Tabelle aus. Was stimmt für welche App? (Ja / Nein / Mit Einschränkung)",
             italic=True, color=GREY, space_after=8)

    headers = ["Aussage", "WhatsApp", "Snapchat"]
    rows_data = [
        ("Meine Inhalte sind Ende-zu-Ende verschlüsselt.",
         "Ja (außer Meta-AI-Chats)", "Nein"),
        ("Wenn ich poste, kann die App damit kommerziell Geld verdienen.",
         "Nein", "Ja, ohne Vergütung"),
        ("Wenn ich meinen Account lösche, sind alle Daten sofort weg.",
         "Nein (Lizenzen bleiben, 30 Tage Speicherung)", "Nein („auf unbestimmte Zeit\")"),
        ("Mein Standort wird automatisch geteilt.",
         "Nein", "Auf Snap Map ja, außer Geistmodus"),
        ("Ich kann das Unternehmen vor einem deutschen Gericht verklagen.",
         "Ja (EU-Recht)", "Nein (Schiedsverfahren USA)"),
        ("Daten gehen an Strafverfolgungsbehörden, wenn diese fragen.",
         "Metadaten ja, Inhalte nein (E2E)", "Daten und Inhalte möglich"),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    widths = [Cm(7.5), Cm(4.5), Cm(4.5)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    # Header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        c.text = ""
        para = c.paragraphs[0]
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, "1F3A4D" if i == 0 else ("169C4C" if i == 1 else "C9A800"))

    # Rows
    for ri, (frage, wa, snap) in enumerate(rows_data, start=1):
        c0 = table.rows[ri].cells[0]
        c0.text = ""
        c0.width = widths[0]
        p = c0.paragraphs[0]
        r = p.add_run(frage)
        r.font.size = Pt(10)

        for ci, val in enumerate([wa, snap], start=1):
            c = table.rows[ri].cells[ci]
            c.text = ""
            c.width = widths[ci]
            p = c.paragraphs[0]
            if loesungen:
                r = p.add_run(val)
                r.font.size = Pt(9.5)
                r.font.color.rgb = DARK
            else:
                r = p.add_run(" ")
                r.font.size = Pt(10)

    # ──────────────────────────────────────────────────────────
    # AUFGABE 4 — Reflexion
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 4, "Reflexion — deine Meinung zählt",
                   color=SNAP_YELLOW)
    add_para(doc, "Beantworte zwei der folgenden vier Fragen ausführlich (je 3–5 Sätze).",
             italic=True, color=GREY, space_after=8)

    fragen = [
        "1. Würdest du diese AGB unterschreiben, wenn man sie dir auf Papier in der Bank vorlegt? Begründe.",
        "2. Was machst du ab heute anders auf WhatsApp oder Snapchat? Nenne mindestens zwei konkrete Schritte.",
        "3. Sollten Apps für Kinder und Jugendliche unter 16 strenger reguliert oder verboten werden? Pro/Kontra.",
        "4. Was bedeutet es für dich, dass Snap dein Bitmoji „unbefristet kommerziell\" nutzen darf?",
    ]

    if loesungen:
        # Erwartungshorizonte
        erwart = [
            ("1.", "Erwartung: Argumente erkennen — niemand würde 30 Seiten Vertrag unterschreiben, ohne zu lesen. Bewusstsein für die Asymmetrie."),
            ("2.", "Erwartung: Konkrete Schritte wie Geistmodus aktivieren, Meta AI nicht in Gruppen, weniger öffentliche Inhalte, AGB-Zusammenfassungen lesen, E2E-Backup, Adressbuch-Zugriff prüfen."),
            ("3.", "Erwartung: Pro-Argumente (Schutz, Reife, DSGVO) und Kontra-Argumente (Kommunikation, Teilhabe, Gleichstellung). Keine richtige Antwort — Bewertung der Argumentationstiefe."),
            ("4.", "Erwartung: Verständnis, dass das digitale Abbild („Bitmoji = ich\") rechtlich an ein US-Unternehmen verschenkt wird. Persönlichkeitsrechte vs. Vertragsfreiheit."),
        ]
        for nr, txt in erwart:
            add_para(doc, txt, size=10, color=GREY, italic=True, space_after=8)
    else:
        for f in fragen:
            add_para(doc, f, size=11, bold=True, color=DARK, space_after=4)
            add_lines(doc, 4)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ──────────────────────────────────────────────────────────
    # FOOTER
    # ──────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Quellen: WhatsApp-Nutzungsbedingungen · Snap Servicebestimmungen · Verbraucherzentrale.de · LDI NRW · ecoprotec.de")
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = GREY

    return doc


if __name__ == "__main__":
    out_dir = "/Users/marcolemke/agb-aufklaerung"
    ab = baue_arbeitsblatt(loesungen=False)
    ab_path = f"{out_dir}/Arbeitsblatt_AGB.docx"
    ab.save(ab_path)
    print(f"✓ Arbeitsblatt gespeichert: {ab_path}")

    lo = baue_arbeitsblatt(loesungen=True)
    lo_path = f"{out_dir}/Arbeitsblatt_AGB_Loesungen.docx"
    lo.save(lo_path)
    print(f"✓ Lösungsblatt gespeichert: {lo_path}")
